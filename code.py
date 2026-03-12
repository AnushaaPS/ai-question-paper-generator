import io
import re
import time
import random
import zipfile
import pdfplumber
import docx2txt
import streamlit as st
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import requests

# -----------------------
# STREAMLIT SETUP
# -----------------------
st.set_page_config(layout="wide")
st.title("KCET - AI Question Paper Generator")


# -----------------------
# HUGGINGFACE TOKEN + MODEL INIT
# -----------------------
hf_token = st.text_input("Enter HuggingFace Access Token", type="password")
model_name = st.text_input("HuggingFace Model ID", 
                           value="meta-llama/Meta-Llama-3-8B-Instruct")

if not hf_token:
    st.warning("Enter HuggingFace Access Token to continue.")
    st.stop()

API_URL = "https://router.huggingface.co/v1/chat/completions"

def hf_generate(prompt):
    payload = {
        "model": model_name, 
        "messages": [
            {"role": "user", "content": prompt}
        ],
        "temperature": 0.8,
        "top_p": 0.9
        # no max_tokens → HF automatically handles length
    }

    try:
        response = requests.post(
            API_URL,
            headers={
                "Authorization": f"Bearer {hf_token}",
                "Content-Type": "application/json"
            },
            json=payload
        )

        if response.status_code != 200:
            return f"⚠ HF API Error {response.status_code}: {response.text}"

        data = response.json()

        # Extract generated text from OpenAI-style response
        return data["choices"][0]["message"]["content"]

    except Exception as e:
        return f"⚠ Error: {e}"

# Use HF generator for both functions
def gen_flash(prompt):
    return hf_generate(prompt)

def gen_pro(prompt):
    return hf_generate(prompt)

# -----------------------
# GENERATION CONFIG
# -----------------------
GEN_CFG = {
    "temperature": 0.8,
    "top_p": 0.9,
    "top_k": 40
}

def normalize_question(text: str) -> str:
    text = text.lower()
    text = re.sub(r"\(\s*\d+\s*marks?\s*\)", "", text, flags=re.IGNORECASE)
    text = re.sub(r"[^a-z0-9]+", " ", text)
    text = re.sub(r"\s+", " ", text).strip()
    return text

def is_too_similar(q1: str, q2: str, threshold: float = 0.75) -> bool:
    n1 = normalize_question(q1)
    n2 = normalize_question(q2)
    if not n1 or not n2:
        return False
    set1 = set(n1.split())
    set2 = set(n2.split())
    if not set1 or not set2:
        return False
    overlap = len(set1 & set2) / max(len(set1), len(set2))
    return overlap >= threshold

# ---------------------------
# SYLLABUS EXTRACTION & DOCX
# ---------------------------
def extract_pdf(file):
    txt = ""
    try:
        with pdfplumber.open(file) as pdf:
            for p in pdf.pages:
                t = p.extract_text()
                if t:
                    txt += t + "\n"
    except:
        pass
    return txt.strip()

def extract_docx(file):
    try:
        return docx2txt.process(file).strip()
    except:
        return ""

def extract_bold_units(file):
    try:
        doc = Document(file)
    except:
        return []
    out = []
    for p in doc.paragraphs:
        found = None
        for r in p.runs:
            if r.bold and r.text and r.text.strip():
                found = r.text.strip()
                break
        if not found:
            text = p.text.strip()
            if text and text.upper() == text and 2 <= len(text.split()) <= 12:
                found = text
        if found:
            if not out or (found not in out):
                out.append(found)
    return out

def auto_units(text):
    out = []
    for line in text.split("\n"):
        line = line.strip()
        if not line:
            continue
        if line.upper() == line and 2 <= len(line.split()) <= 12:
            if line not in out:
                out.append(line)
    return out[:50]

def extract_units_with_content(text):
    units = {}
    current_unit = None

    for line in text.split("\n"):
        line = line.strip()
        if not line:
            continue

        if line.isupper() and len(line.split()) <= 6:
            current_unit = line.strip().upper()
            units[current_unit] = []
        elif current_unit:
            units[current_unit].append(line)

    for k in units:
        units[k] = " ".join(units[k])

    return units

def split_unit_semantically(unit_text):
    sentences = re.split(r'[.;]', unit_text)
    sentences = [s.strip() for s in sentences if s.strip()]

    n = len(sentences)
    mid = n // 2

    return {
        "First Half": " ".join(sentences[:mid]),
        "Second Half": " ".join(sentences[mid:])
    }

def get_unit_portion(parts, portion):
    n = len(parts)
    if n == 0:
        return ""

    mid = n // 2
    if portion == "First Half":
        return " | ".join(parts[:mid])
    else:
        return " | ".join(parts[mid:])

def split_unit_quarters(unit_text):
    sentences = re.split(r'[.;]', unit_text)
    sentences = [s.strip() for s in sentences if s.strip()]

    n = len(sentences)
    q = max(1, n // 4)

    return {
        1: " ".join(sentences[:q]),
        2: " ".join(sentences[q:2*q]),
        3: " ".join(sentences[2*q:3*q]),
        4: " ".join(sentences[3*q:])
    }

def format_units_as_request_b(units_raw):
    formatted = []
    for idx, u in enumerate(units_raw):
        name = u.strip()
        name = re.sub(r"^[\-\:\s]+|[\-\:\s]+$", "", name)
        name = re.sub(r"\s+", " ", name).strip().upper()
        formatted.append(f"UNIT {idx+1}: {name}")
    return formatted

def make_docx(content, logo_bytes, header):
    doc = Document()
    if logo_bytes:
        try:
            logo_bytes.seek(0)
            doc.add_picture(logo_bytes, width=Inches(2))
            doc.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        except:
            pass
    doc.add_heading(header["college"], level=1)
    doc.add_paragraph(f"{header['code']} — {header['name']}")
    doc.add_paragraph(f"Marks: {header['marks']}   Duration: {header['duration']}")
    doc.add_paragraph("")
    for line in content.split("\n"):
        doc.add_paragraph(line)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

# ---------------------------
# PROMPT BUILDER
# ---------------------------
def build_prompt_question(syllabus, sec, q_number, q_marks, q_type, klevel, pattern=None):
    common = f"""
STRICT ENFORCEMENT:
- If content scope is provided, you MUST generate the question strictly within it
- Violating scope will be considered incorrect

Generate exactly ONE question for {sec}.

SYLLABUS:
{syllabus}

RULES:
- Question number: {q_number}
- Marks: {q_marks}
- Type: {q_type}
- K-Level: {klevel}
- You MUST generate the question only according to the K-Level specified above.
- DO NOT mention units, do not print which unit the question is from.

⭐ IMPORTANT NEW CONDITIONS (must follow strictly):
1. DO NOT repeat any previously generated question.  
   - The question must be completely new and unique.  
   - Avoid same structure, wording, or meaning.

2. For EITHER–OR questions:
   - Both (a) and (b) MUST come from the SAME ASSIGNED UNIT ONLY.
   - They must NOT overlap or be similar.
"""

    # Combined K-level handling
    if "-" in klevel:
        levels = klevel.split("-")
        common += f"""
- This is a combined K-level question including: {klevel}.
- The question MUST require thinking skills from ALL the K-levels mentioned.
- Include components for each level such as:
  * K2: Explanation / description
  * K3: Application-based mini problem or example
  * K4: Analytical comparison, reasoning or inference
"""
    else:
        common += f"""
- You MUST generate the question only according to the single K-Level specified and must NOT mix other levels.
"""

    # MCQ rules
    if q_type == "MCQ":
        common += """
IMPORTANT:
- THIS QUESTION MUST BE AN MCQ.
- Format strictly:
  Question text?
  A) Option 1
  B) Option 2
  C) Option 3
  D) Option 4
- DO NOT write explanations.
- DO NOT write long answers.
- DO NOT produce descriptive content.
"""

    # Short answer rules
    elif q_type == "Short Answer":
        common += """
- Short Answer: 1–3 lines.
- Keep question direct and focused.
"""

    # Long answer rules
    elif q_type == "Long Answer":
        if pattern and "+" in str(pattern):
            parts = [p.strip() for p in str(pattern).split("+") if p.strip().isdigit()]
            if len(parts) == 2:
                common += f"""
- Long Answer MUST include exactly two subparts:
  (a) ... ({parts[0]} marks)
  (b) ... ({parts[1]} marks)
- Subparts must be meaningfully related.
"""
            else:
                common += f"""
- Long Answer must include subparts summing to {q_marks} marks.
"""
        else:
            common += f"""
- Long Answer: one full question worth {q_marks} marks (NO subparts).
"""

    # Final output rule
    common += "\nOutput ONLY the question text (the app will add numbering).\n"
    return common

def build_answer_prompt(syllabus, sec, q_number, question_text, q_type):
    """
    Single-question answer prompt builder. Uses the K-level implicitly because question_text was generated using it.
    """
    return f"""
Generate an ANSWER for Question {q_number} in {sec}.

SYLLABUS:
{syllabus}

QUESTION:
{question_text}

RULES:
- Output ONLY the answer (do not repeat the question).
- For MCQ: give the correct option (A-D) and a one-line reason.
- For Short Answer: 2-3 lines.
- For Long Answer: give key points and subpart answers (if question has subparts, provide separate answers for (a) and (b)).
- Keep concise.
"""

# ---------------------------
# UI — UPLOAD
# ---------------------------
st.header("Upload Syllabus & Logo")
logo_file = st.file_uploader("Logo", ["png", "jpg", "jpeg"])
syll_file = st.file_uploader("Syllabus PDF/DOCX", ["pdf", "docx"])

syllabus = ""
units_detected_raw = []
units_detected_formatted = []

if syll_file:
    ext = syll_file.name.split(".")[-1].lower()
    if ext == "docx":
        syllabus = extract_docx(syll_file)
        units_detected_raw = extract_bold_units(syll_file)
        if not units_detected_raw:
            units_detected_raw = auto_units(syllabus)
    else:
        syllabus = extract_pdf(syll_file)
        units_detected_raw = auto_units(syllabus)

    units_detected_formatted = format_units_as_request_b(units_detected_raw)

st.subheader("Detected Units")
if units_detected_formatted:
    st.text_area("Detected Units (editable)", "\n".join(units_detected_formatted), height=150, key="detected_units_area")
    st.success(f"{len(units_detected_formatted)} units detected.")
else:
    st.text_area("Detected Units (editable)", "", height=150, key="detected_units_area")
    st.info("0 units detected.")

unit_content_map = extract_units_with_content(syllabus)

edited = st.session_state.get("detected_units_area", "").split("\n")
units_list = []
for line in edited:
    line = line.strip()
    if not line:
        continue
    m = re.match(r'UNIT\s*\d+\s*[:\-]\s*(.*)', line, flags=re.IGNORECASE)
    if m:
        name = m.group(1).strip().upper()
    else:
        name = line.upper()
    units_list.append(name)
units_list = [f"UNIT {i+1}: {units_list[i]}" for i in range(len(units_list))]
num_units = len(units_list)

# ---------------------------
# EXAM CONFIG
# ---------------------------
st.header("Exam Configuration")
college = st.text_input("College Name")
code = st.text_input("Course Code")
name = st.text_input("Course Name")
marks_total = st.number_input("Total Marks", 1, 500, 50)
duration = st.selectbox("Duration", ["45 Minutes", "1 Hour", "1.5 Hours", "2 Hours", "3 Hours"])
sections = st.multiselect("Sections", ["PART A", "PART B", "PART C"], default=["PART A"])
set_count = st.number_input("Number of Sets", 1, 10, 1)

# ---------------------------
# SECTION CONFIG
# ---------------------------
st.header("Section Configurations")
section_cfg = {}
computed_marks = 0
K_LEVEL_OPTIONS = ["K1 - Remember", "K2 - Understand", "K3 - Apply",
                   "K4 - Analyze", "K5 - Evaluate", "K6 - Create", "K2-K3", "K2-K4", "K3-K4"]
PARTA_TYPE_OPTIONS = ["MCQ", "Short Answer", "Long Answer"]

for sec in sections:
    st.subheader(sec)
    tq = st.number_input(f"{sec} - Total Questions", 1, 100, 5, key=f"{sec}_tq")
    mq = st.number_input(f"{sec} - Marks/Question", 1, 200, 2, key=f"{sec}_mq")

    qtype_section = None
    if sec != "PART A":
        qtype_section = "Long Answer"

    st.markdown("### Unit Distribution")
    auto = tq // num_units if num_units else 1
    rem = tq % num_units if num_units else 0
    dist = {}
    remain = tq
    for i in range(num_units):
        default = auto + (1 if i < rem else 0)
        default = min(default, remain)
        val = st.number_input(f"{sec}: Unit {i+1}", 0, remain, default, key=f"{sec}unit{i}")
        dist[f"Unit {i+1}"] = val
        remain -= val

    st.markdown("### K-Level for each Question")
    klevels_list = []
    for q in range(1, tq + 1):
        kl = st.selectbox(f"{sec} - Q{q}", K_LEVEL_OPTIONS, key=f"{sec}klevel{q}")
        klevels_list.append(kl)

    q_formats = []
    unit_portions = []

    if sec in ["PART B", "PART C"]:
        st.markdown("### Long Answer Pattern")
        for q in range(1, tq + 1):
            fmt = st.selectbox(
                f"{sec} - Q{q} Format",
                ["Single", "Split-up"],
                key=f"{sec}format{q}"
            )
            q_formats.append(fmt)

    if sec in ["PART B", "PART C"]:
        st.markdown("### Unit Portion Selection (Long Answer)")
        for q in range(1, tq + 1):
            portion = st.selectbox(
                f"{sec} - Q{q} Unit Portion",
                ["First Half", "Second Half"],
                key=f"{sec}_portion_{q}"
            )
            unit_portions.append(portion)
    else:
        q_formats = ["Single"] * tq
        unit_portions = [None] * tq


    parta_qtypes = None
    if sec == "PART A":
        st.markdown("### PART A Question Type")
        parta_qtypes = []
        for q in range(1, tq + 1):
            qt = st.selectbox(f"{sec} - Q{q} Type", PARTA_TYPE_OPTIONS, key=f"{sec}qt{q}")
            parta_qtypes.append(qt)

    computed_marks += tq * mq

    section_cfg[sec] = {
        "total_questions": tq,
        "marks": mq,
        "type": qtype_section,
        "klevels": klevels_list,
        "qtypes_parta": parta_qtypes,
        "distribution": dist,
        "formats": q_formats,
        "unit_portions": unit_portions
    }

st.info(f"Computed Total Marks = {computed_marks} | Entered = {marks_total}")

# ---------------------------
# AUTO PATTERN FOR PART B
# ---------------------------
def auto_pattern(marks):
    if marks == 16:
        return random.choice(["8+8", "10+6", "12+4"])
    return str(marks)

# ---------------------------
# HEADER BUILDER
# ---------------------------
def build_header():
    return {
        "college": college,
        "code": code,
        "name": name,
        "marks": marks_total,
        "duration": duration
    }

# ---------------------------
# GENERATE BUTTON
# ---------------------------
if st.button("Generate Question Papers"):
    if not syllabus:
        st.error("Upload syllabus first.")
        st.stop()
    for sec in sections:
        expected = section_cfg[sec]["total_questions"]
        actual = sum(section_cfg[sec]["distribution"].values())
        if expected != actual:
            st.error(f"Unit mismatch in {sec}: expected {expected}, got {actual}")
            st.stop()
    if computed_marks != marks_total:
        st.error("Mismatch in total marks")
        st.stop()

    st.success("Generating...")

    logo_bytes = None
    if logo_file:
        logo_bytes = io.BytesIO(logo_file.read())

    header = build_header()
    zip_buf = io.BytesIO()
    zipf = zipfile.ZipFile(zip_buf, "w", zipfile.ZIP_DEFLATED)

    global_seen_questions = []

    for s in range(int(set_count)):
        st.header(f"SET {s+1}")
        qp_sections = []
        seen_questions_this_set = []
        global_qnum = 1

        for sec in sections:
            cfg = section_cfg[sec]
            tq = cfg["total_questions"]
            q_marks = cfg["marks"]

            pattern = None
            if sec == "PART B":
                pattern = auto_pattern(q_marks)

            q_texts = []
            st.info(f"Generating {sec} ({tq} questions (Set {s+1})")

            unit_assignment = []
            for ui in range(num_units):
                count = cfg["distribution"].get(f"Unit {ui+1}", 0)
                unit_assignment.extend([ui+1] * count)
            if len(unit_assignment) != tq:
                unit_assignment = [((i % num_units) + 1) for i in range(tq)]

            for i in range(tq):
                qnum = global_qnum
                global_qnum += 1
                klevel = cfg["klevels"][i]

                if sec == "PART A":
                    q_type = cfg["qtypes_parta"][i]
                else:
                    q_type = cfg["type"]

                unit_for_q = unit_assignment[i]
                unit_label = units_list[unit_for_q - 1]

                try:
                    m = re.match(r'UNIT\s*\d+\s*[:\-]\s*(.*)', unit_label, flags=re.IGNORECASE)
                    unit_name = (m.group(1) if m else unit_label).strip().upper()

                    portion = cfg["unit_portions"][i]

                    unit_text = unit_content_map.get(unit_name, "")

                    if not unit_text:
                        unit_text = syllabus   # safe fallback ONLY if extraction fails

                    unit_parts = split_unit_semantically(unit_text)
                    selected_unit_text = unit_parts.get(portion, unit_text)

                    unit_snippet = f"""
                    You MUST generate the question ONLY from the following syllabus portion:

                    {selected_unit_text}

                    STRICT RULES:
                    - Do NOT use topics outside this portion
                    - Do NOT combine multiple unit sections
                    """

                except:
                    unit_text = syllabus 
                    unit_snippet = f"Full Syllabus:\n{syllabus}"

                q_prompt = build_prompt_question(unit_snippet, sec, qnum, q_marks, q_type, klevel, pattern)

                # ---------- EITHER-OR LONG ANSWER FEATURE ----------
                if q_type == "Long Answer" and sec in ["PART B", "PART C"]:
                    q_format = cfg["formats"][i]
                    # SPLIT EITHER-OR FORMAT (i & ii)
                    if q_format == "Split-up":
                        chosen_pattern = random.choice(["8+8", "10+6", "12+4"])
                        p1, p2 = [int(x) for x in chosen_pattern.split("+")]
                        quarters = split_unit_quarters(unit_text)

                        if portion == "First Half":
                            part_i = quarters[1]
                            part_ii = quarters[2]
                        else:
                            part_i = quarters[3]
                            part_ii = quarters[4]

                        q_a_i = gen_flash(build_prompt_question(f"Use ONLY this content:\n{part_i}", sec, f"{qnum} a i", p1, q_type, klevel))
                        q_a_ii = gen_flash(build_prompt_question(f"Use ONLY this content:\n{part_ii}", sec, f"{qnum} a ii", p2, q_type, klevel))
                        q_b_i = gen_flash(build_prompt_question(f"Use ONLY this content:\n{part_i}", sec, f"{qnum} b i", p1, q_type, klevel))
                        q_b_ii = gen_flash(build_prompt_question(f"Use ONLY this content:\n{part_ii}", sec, f"{qnum} b ii", p2, q_type, klevel))

                        q_texts.append(
                            f"{qnum}) a) i) {q_a_i} ({p1} Marks)\n"
                            f"{qnum}) a) ii) {q_a_ii} ({p2} Marks)\n\n"
                            f"   OR\n\n"
                            f"{qnum}) b) i) {q_b_i} ({p1} Marks)\n"
                            f"{qnum}) b) ii) {q_b_ii} ({p2} Marks)"
                        )
                        continue
                    # SINGLE LONG QUESTION EITHER-OR format
                    else:
                        q_a = gen_flash(build_prompt_question(unit_snippet, sec, f"{qnum} a", q_marks, q_type, klevel))
                        q_b = gen_flash(build_prompt_question(unit_snippet, sec, f"{qnum} b", q_marks, q_type, klevel))

                        q_texts.append(
                            f"{qnum}) a) {q_a}\n\n"
                            f"   OR\n\n"
                            f"{qnum}) b) {q_b}"
                        )
                        continue
                    
                best_q_text = None
                q_text_raw = ""
                generated = "" 

                for attempt in range(4):
                    generated = gen_flash(q_prompt)

                    if not generated or generated.startswith("⚠"):
                        continue

                    q_text_raw = re.sub(r"^\s*(?:Q|Question)?\s*\d+[:\.\)]\s*", "", generated.strip())
                    duplicate_found = False
                    for prev in seen_questions_this_set + global_seen_questions:
                        if is_too_similar(q_text_raw, prev):
                            duplicate_found = True
                            break
                    if not duplicate_found:
                        best_q_text = q_text_raw
                        break

                if best_q_text is None or "⚠" in best_q_text or "Error" in best_q_text:
                    best_q_text = "⚠ Could not generate question"

                seen_questions_this_set.append(best_q_text)
                global_seen_questions.append(best_q_text)

                q_texts.append(f"{qnum}. {best_q_text} ({q_marks} Marks)")

            section_q_block = f"{sec}\n{'-'*40}\n" + "\n\n".join(q_texts)
            qp_sections.append(section_q_block)

        full_qp = "\n\n".join(qp_sections)
        qp_doc = make_docx(full_qp, logo_bytes, header)

        st.subheader(f"Preview — SET {s+1}")
        st.text_area(f"QP_{s+1}", full_qp, height=300)
        st.download_button(f"Download QP Set {s+1}", qp_doc.getvalue(), f"QP_Set{s+1}.docx")

        zipf.writestr(f"QP_Set{s+1}.docx", qp_doc.getvalue())

    zipf.close()
    zip_buf.seek(0)
    st.download_button("Download All Sets (ZIP)", zip_buf.getvalue(), "All_Sets.zip")

    st.success("Completed!")
